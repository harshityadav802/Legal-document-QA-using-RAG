import os
from pathlib import Path
from typing import List, Optional

from langchain_core.documents import Document

from src.segmentation.legal_segmenter import LegalChunk, segment_document
from src.segmentation.preprocessor import preprocess
from src.vectorstore.store import build_vectorstore, load_vectorstore


def _read_file(file_path: str) -> str:

    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = path.suffix.lower()

    if suffix == ".txt":
        return path.read_text(encoding="utf-8")

    if suffix == ".pdf":

        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            pages = [page.extract_text() or "" for page in reader.pages]

            return "\n\n".join(pages)

        except ImportError as exc:

            raise ImportError(
                "pypdf required. Install with: pip install pypdf"
            ) from exc

    if suffix in (".docx", ".doc"):

        try:
            import docx

            doc = docx.Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            return "\n\n".join(paragraphs)

        except ImportError as exc:

            raise ImportError(
                "python-docx required. Install with: pip install python-docx"
            ) from exc

    raise ValueError(f"Unsupported file format: {suffix}")


def _chunks_to_documents(chunks: List[LegalChunk]) -> List[Document]:

    docs = []

    for chunk in chunks:

        metadata = {
            "breadcrumb": chunk.breadcrumb,
            "section_type": chunk.section_type,
            "document_name": chunk.document_name,
            "is_operative": chunk.is_operative,
            "cross_references": ", ".join(chunk.cross_references),
            "chunk_index": chunk.chunk_index,
        }

        docs.append(
            Document(
                page_content=chunk.text,
                metadata=metadata,
            )
        )

    return docs


def ingest_document(
    file_path: str,
    document_name: Optional[str] = None,
    store_path: Optional[str] = None,
    append: bool = False,
) -> List[Document]:

    effective_store_path = store_path or os.getenv(
        "VECTORSTORE_PATH", "data/vectorstore"
    )

    if document_name is None:
        document_name = Path(file_path).stem.replace("_", " ").title()

    raw_text = _read_file(file_path)

    cleaned_text, doc_type = preprocess(raw_text)
    print(f"[ingest] Document type: {doc_type}")

    chunks = segment_document(cleaned_text, document_name=document_name)
    print(f"[ingest] Created {len(chunks)} chunks")

    documents = _chunks_to_documents(chunks)

    if append:

        try:

            existing = load_vectorstore(effective_store_path)

            existing.add_documents(documents)
            existing.save_local(effective_store_path)

            print(f"[ingest] Added {len(documents)} documents")

        except FileNotFoundError:

            build_vectorstore(documents, store_path=effective_store_path)

            print(f"[ingest] Created new vector store")

    else:

        build_vectorstore(documents, store_path=effective_store_path)

        print(f"[ingest] Built vector store with {len(documents)} docs")

    return documents
def ingest_dataset(
    dataset_path: str,
    store_path: Optional[str] = None,
):

    dataset = Path(dataset_path)

    if not dataset.exists():
        raise FileNotFoundError(f"Dataset folder not found: {dataset_path}")

    pdf_files = list(dataset.rglob("*.pdf"))

    print(f"Found {len(pdf_files)} PDF files")

    for i, pdf_file in enumerate(pdf_files):

        print(f"\n[{i+1}/{len(pdf_files)}] Processing: {pdf_file.name}")

        ingest_document(
            file_path=str(pdf_file),
            document_name=pdf_file.stem,
            store_path=store_path,
            append=True,
        )
