"""Document ingestion pipeline.

Reads a document file, preprocesses it, segments it into legal chunks,
embeds each chunk, and stores them in a FAISS vector store.
"""

import os
from pathlib import Path
from typing import List, Optional

from langchain_core.documents import Document

from src.segmentation.legal_segmenter import LegalChunk, segment_document
from src.segmentation.preprocessor import preprocess
from src.vectorstore.store import build_vectorstore, load_vectorstore


def _read_file(file_path: str) -> str:
    """Read text from a .txt, .pdf, or .docx file."""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    suffix = path.suffix.lower()

    if suffix == ".txt":
        return path.read_text(encoding="utf-8")

    if suffix == ".pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n\n".join(pages)
        except ImportError as exc:
            raise ImportError(
                "pypdf is required for PDF support. Install with: pip install pypdf"
            ) from exc

    if suffix in (".docx", ".doc"):
        try:
            import docx

            doc = docx.Document(str(path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        except ImportError as exc:
            raise ImportError(
                "python-docx is required for Word support. "
                "Install with: pip install python-docx"
            ) from exc

    raise ValueError(f"Unsupported file format: {suffix}")


def _chunks_to_documents(chunks: List[LegalChunk]) -> List[Document]:
    """Convert LegalChunk objects to LangChain Document objects."""
    docs = []
    for chunk in chunks:
        metadata = {
            "breadcrumb": chunk.breadcrumb,
            "section_type": chunk.section_type,
            "document_name": chunk.document_name,
            "is_operative": chunk.is_operative,
            "cross_references": ", ".join(chunk.cross_references),
            "chunk_index": chunk.chunk_index,
        }
        docs.append(Document(page_content=chunk.text, metadata=metadata))
    return docs


def ingest_document(
    file_path: str,
    document_name: Optional[str] = None,
    store_path: Optional[str] = None,
    append: bool = False,
) -> List[Document]:
    """Ingest a legal document into the vector store.

    Args:
        file_path:     Path to the document file (.txt, .pdf, or .docx).
        document_name: Human-readable name (defaults to filename).
        store_path:    Directory to persist the FAISS index.
        append:        If True, merge into an existing store; otherwise replace.

    Returns:
        List of LangChain Documents that were ingested.
    """
    effective_store_path = store_path or os.getenv(
        "VECTORSTORE_PATH", "data/vectorstore"
    )
    if document_name is None:
        document_name = Path(file_path).stem.replace("_", " ").title()

    # 1. Read raw text
    raw_text = _read_file(file_path)

    # 2. Preprocess
    cleaned_text, doc_type = preprocess(raw_text)
    print(f"[ingest] Document type detected: {doc_type}")

    # 3. Segment
    chunks = segment_document(cleaned_text, document_name=document_name)
    print(f"[ingest] Segmented into {len(chunks)} chunks")

    # 4. Convert to LangChain Documents
    documents = _chunks_to_documents(chunks)

    # 5. Build or append to vector store
    if append:
        try:
            existing = load_vectorstore(effective_store_path)
            existing.add_documents(documents)
            existing.save_local(effective_store_path)
            print(f"[ingest] Appended {len(documents)} docs to existing store")
        except FileNotFoundError:
            build_vectorstore(documents, store_path=effective_store_path)
            print(f"[ingest] Created new store with {len(documents)} docs")
    else:
        build_vectorstore(documents, store_path=effective_store_path)
        print(f"[ingest] Built new store with {len(documents)} docs")

    return documents
